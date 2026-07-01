"""File upload and text extraction utility functions.

Provides helper functions for:
- Validating file extensions, MIME types, and sizes.
- Preventing directory traversal attacks.
- Saving uploaded files with unique names.
- Extracting plain text from PDF using PyMuPDF and DOCX using python-docx.
"""

import os
import uuid
import re
from pathlib import Path
from fastapi import UploadFile

from app.config import get_settings
from app.core.exceptions import ValidationException
from app.core.logging import get_logger

logger = get_logger(__name__)

ALLOWED_EXTENSIONS = {".pdf", ".docx"}
ALLOWED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5 Megabytes


def ensure_upload_directory() -> Path:
    """Create the upload directory if it does not exist.

    Returns:
        Path to the upload directory.
    """
    settings = get_settings()
    upload_path = Path(settings.UPLOAD_DIRECTORY)
    upload_path.mkdir(parents=True, exist_ok=True)
    return upload_path


def validate_file(file: UploadFile) -> None:
    """Validate uploaded file extension, MIME type, and size constraints.

    Args:
        file: The UploadFile to validate.

    Raises:
        ValidationException: If the file is invalid.
    """
    filename = file.filename or ""
    suffix = Path(filename).suffix.lower()

    # 1. Validate Extension
    if suffix not in ALLOWED_EXTENSIONS:
        raise ValidationException("Unsupported file type. Only PDF and DOCX files are allowed.")

    # 2. Validate MIME Type
    mime_type = file.content_type or ""
    if mime_type not in ALLOWED_MIME_TYPES:
        raise ValidationException("Invalid MIME type. Only PDF and DOCX documents are allowed.")


async def save_upload_file(file: UploadFile) -> Path:
    """Save an uploaded file securely under uploads/ directory with a unique name.

    Ensures no directory traversal and size limits are enforced.

    Args:
        file: The UploadFile instance.

    Returns:
        The Path to the saved file.
    """
    # Validate extension and mime type first
    validate_file(file)

    upload_dir = ensure_upload_directory().resolve()
    
    file_id = uuid.uuid4()
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        suffix = ".pdf"  # Fallback

    unique_filename = f"{file_id}{suffix}"
    dest_path = (upload_dir / unique_filename).resolve()

    # Prevent directory traversal attacks
    if not dest_path.is_relative_to(upload_dir):
        raise ValidationException("Invalid destination file path.")

    size_counter = 0
    try:
        with open(dest_path, "wb") as buffer:
            while chunk := await file.read(1024 * 1024):
                size_counter += len(chunk)
                if size_counter > MAX_FILE_SIZE:
                    # Clean up file and raise error
                    buffer.close()
                    dest_path.unlink(missing_ok=True)
                    raise ValidationException("File size exceeds the 5MB upload limit.")
                buffer.write(chunk)
        
        logger.info("Saved upload file to: %s (size: %d bytes)", dest_path.name, size_counter)
        return dest_path
    except ValidationException:
        raise
    except Exception as exc:
        logger.error("Failed to save uploaded file: %s", exc)
        raise ValidationException("Failed to save uploaded file.") from exc


def extract_text_from_pdf(pdf_path: Path) -> str:
    """Extract and normalize plain text from a PDF file using PyMuPDF (fitz).

    Args:
        pdf_path: Path to the PDF file.

    Returns:
        The clean extracted text string.
    """
    if not pdf_path.exists():
        raise ValidationException("PDF file does not exist.")

    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(pdf_path))
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()

        full_text = "\n".join(text_parts)
        return clean_extracted_text(full_text)
    except Exception as exc:
        logger.error("Failed to extract text from PDF using PyMuPDF: %s", exc)
        raise ValidationException("Failed to extract text from the PDF file.") from exc


def extract_text_from_docx(docx_path: Path) -> str:
    """Extract and normalize plain text from a DOCX file using python-docx.

    Args:
        docx_path: Path to the DOCX file.

    Returns:
        The clean extracted text string.
    """
    if not docx_path.exists():
        raise ValidationException("DOCX file does not exist.")

    try:
        import docx
        doc = docx.Document(str(docx_path))
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text_parts.append(paragraph.text)
        
        # Also extract table text if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text_parts.append(cell.text)

        full_text = "\n".join(text_parts)
        return clean_extracted_text(full_text)
    except Exception as exc:
        logger.error("Failed to extract text from DOCX using python-docx: %s", exc)
        raise ValidationException("Failed to extract text from the DOCX file.") from exc


def clean_extracted_text(text: str) -> str:
    """Clean extra whitespaces, normalize encodings, and return clean text."""
    # Normalize unicode encoding
    import unicodedata
    normalized = unicodedata.normalize("NFKC", text)

    # Clean consecutive spaces and carriage returns
    cleaned = re.sub(r"[ \t]+", " ", normalized)
    cleaned = re.sub(r"\n\s*\n+", "\n\n", cleaned)
    return cleaned.strip()
